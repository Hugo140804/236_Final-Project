# -*- coding: utf-8 -*-
"""Generate Diagram-Blockchain-API.docx berisi Use Case, ERD, dan User Flow diagram."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAM = os.path.join(BASE, 'laporan', 'diagrams')
OUTPUT = os.path.join(BASE, 'laporan', 'Diagram-Blockchain-API.docx')

doc = Document()

# ---------- Konfigurasi default ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

for level, size, color in [(0, 20, '1F3864'), (1, 16, '2E5395'), (2, 13, '2E5395')]:
    style = doc.styles['Heading %d' % (level + 1)]
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.name = 'Calibri'

# ---------- Halaman judul ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('LAPORAN FINAL PROJECT 236')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor.from_string('1F3864')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Diagram Sistem — Blockchain API\n(SaaS: Data via API Key, Autentikasi JWT)')
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string('2E5395')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('\nNama    : (isi nama Anda)\nNIM      : (isi NIM Anda)\nKelas    : 236\n\n').font.size = Pt(12)

doc.add_paragraph('Dokumen ini berisi tiga diagram utama sistem:\n'
                  '1. Use Case Diagram\n'
                  '2. Entity Relationship Diagram (ERD)\n'
                  '3. Activity Diagram / User Flow').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string('595959')


def add_body(text):
    doc.add_paragraph(text)


def add_bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


# ============================================================
# 1. USE CASE DIAGRAM
# ============================================================
doc.add_heading('1. USE CASE DIAGRAM', level=1)
add_body('Use case diagram menggambarkan interaksi antara aktor dengan sistem '
         'Blockchain API. Terdapat dua aktor:')
add_bullets([
    'Pengembang — mengelola akun, API key, dan data (register, login, kelola API key, kelola blockchain & kategori).',
    'Konsumen API — pihak luar yang hanya membaca data menggunakan API key.',
])
doc.add_picture(os.path.join(DIAGRAM, 'usecase.png'), width=Inches(6.3))
add_caption('Gambar 1. Use Case Diagram Blockchain API')
doc.add_paragraph()
add_bullets([
    'UC-1 Registrasi Akun (Pengembang)',
    'UC-2 Login JWT (Pengembang)',
    'UC-3 Kelola API Key — buat / lihat / nonaktifkan (Pengembang)',
    'UC-4 Kelola Data Blockchain — tambah / ubah / hapus (Pengembang)',
    'UC-5 Kelola Kategori — tambah / ubah / hapus (Pengembang)',
    'UC-6 Lihat Data Blockchain & Kategori (Konsumen API / Pengembang)',
])
doc.add_page_break()

# ============================================================
# 2. ERD
# ============================================================
doc.add_heading('2. ENTITY RELATIONSHIP DIAGRAM (ERD)', level=1)
add_body('ERD menggambarkan 5 tabel database beserta relasinya. Terdapat relasi '
         'one-to-many (pengembang → blockchain, pengembang → api_keys) dan '
         'many-to-many (blockchain ↔ kategori) melalui tabel join BlockchainKategori.')
doc.add_picture(os.path.join(DIAGRAM, 'erd.png'), width=Inches(6.3))
add_caption('Gambar 2. Entity Relationship Diagram Blockchain API')
doc.add_paragraph()
add_bullets([
    'pengembang (id, nama, email, password)',
    'api_keys (id, pengembang_id FK, nama, key, aktif, terakhir_dipakai)',
    'blockchain (id, nama, simbol, deskripsi, tahun_rilis, pengembang_id FK)',
    'kategori (id, nama, deskripsi)',
    'BlockchainKategori (blockchain_id FK, kategori_id FK) — tabel relasi',
])
doc.add_page_break()

# ============================================================
# 3. ACTIVITY DIAGRAM / USER FLOW
# ============================================================
doc.add_heading('3. ACTIVITY DIAGRAM / USER FLOW', level=1)
add_body('Activity diagram berikut menggambarkan alur konsumen dalam mengakses '
         'data menggunakan API key, mulai dari registrasi hingga menerima data.')
doc.add_picture(os.path.join(DIAGRAM, 'activity-userflow.png'), width=Inches(4.3))
add_caption('Gambar 3. Activity Diagram — User Flow Akses Data dengan API Key')
doc.add_paragraph()
add_body('Diagram kedua menggambarkan alur pengembang dalam mengelola data '
         '(operasi tulis) yang dilindungi token JWT.')
doc.add_picture(os.path.join(DIAGRAM, 'activity-crud.png'), width=Inches(4.3))
add_caption('Gambar 4. Activity Diagram — Pengelolaan Data oleh Pengembang (JWT)')

# ---------- Metadata & simpan ----------
props = doc.core_properties
props.title = 'Diagram Sistem - Blockchain API'
props.subject = 'Final Project 236'
props.author = '236 Final Project'
props.keywords = 'use case, erd, activity diagram, blockchain, api'

doc.save(OUTPUT)
print('File Word berhasil dibuat:', OUTPUT)
